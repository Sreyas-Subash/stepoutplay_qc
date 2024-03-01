# pip install python-docx

import sys
import traceback
import docx
import pandas as pd
import os

src_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
parent_dir = os.path.abspath(os.path.join(src_dir, os.pardir))
os.chdir(parent_dir)


class Read():

    read_path = r"read_string\rstring.docx"
    doc_obj = docx.Document(read_path)

    def read(self):
        """
        This function reads the coded string in a docx file to a dataframe
       """

        try:
            # checking for next line character
            if len(self.doc_obj.paragraphs) > 1:
                print(f"At this action {self.doc_obj.paragraphs[1].text.split(',')[0]}, there is an issue. Please Check!")
                print('Terminating Program...')
                sys.exit(0)
            else:
                for paragraph in self.doc_obj.paragraphs:
                    coded_string = paragraph.text
                    coded_string_space_sep = coded_string.split(' ')

                    # checking for space
                    if len(coded_string_space_sep) > 1:
                        print(f"At this action {coded_string_space_sep[1].split(',')[0]}, there is an issue. Please Check!")
                        print('Terminating Program...')
                        sys.exit(0)

                    coded_string_list = coded_string.split(',')
                    df = pd.DataFrame(coded_string_list, columns=['strings'])
                    df[['team', 'jersey_number', 'action', 'notation', 'start_grid', 'end_grid', 'timestamp', 'foot',
                        'special_attribute', 'half']] = df['strings'].str.split('-', expand=True)
                    df = df.drop(columns=['strings'])
                    return df

        except FileNotFoundError:
            print(f"File '{self.read_path}' not found.")
        except Exception as e:
            print(f"An error occurred: {str(e)}")
            traceback.print_exc()

if __name__ == "__main__":
    read_obj = Read()
    df = read_obj.read()

    # print(df.head().to_string())
